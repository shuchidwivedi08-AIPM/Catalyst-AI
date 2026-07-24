import json
from docx import Document
from io import BytesIO
import pytest
from pydantic import ValidationError
from catalyst_ai.ai.artifact_state import build_product_understanding_source_hash, is_product_understanding_stale
from catalyst_ai.ai.response_parser import ArtifactParseError, parse_prd_response, parse_technical_specification_response, parse_user_stories_response
from catalyst_ai.ai.schemas import ArtifactMetadata, ArtifactType, GeneratedArtifact, ProductRequirementsDocument, ProductUnderstanding, TechnicalSpecification, UserStory, UserStoryArtifact
from catalyst_ai.ai.prompts.prd_prompt import build_prd_prompt
from catalyst_ai.ai.prompts.user_stories_prompt import build_user_stories_prompt
from catalyst_ai.ai.prompts.technical_specification_prompt import build_technical_specification_prompt
from catalyst_ai.services.document_export_service import build_artifact_docx


def understanding(): return ProductUnderstanding(executive_summary="Summary", business_problem="Problem", functional_requirements=["Login"])
def metadata(kind): return ArtifactMetadata(artifact_type=kind, stakeholder="Product Owner", context_source="Validated Product Context", product_understanding_source_hash="secret-hash")
def prd(): return ProductRequirementsDocument(title="PRD", executive_summary="Summary", problem_statement="Problem", functional_requirements=["Login"])
def tech(): return TechnicalSpecification(title="Spec", solution_overview="Overview", architecture_summary="Architecture")
def stories(): return UserStoryArtifact(title="Stories", stories=[UserStory(id="US-001", title="Login", persona="User", need="sign in", benefit="access", story="As a User, I want sign in, so that access.", acceptance_criteria=["Given credentials, when submitted, then access is granted."])])

def test_schemas_and_independent_lists():
    assert parse_prd_response(json.dumps(prd().model_dump())).title == "PRD"
    assert parse_user_stories_response(json.dumps(stories().model_dump())).stories[0].id == "US-001"
    assert parse_technical_specification_response(json.dumps(tech().model_dump())).title == "Spec"
    with pytest.raises(ValidationError): ArtifactMetadata(artifact_type="Bad", stakeholder="x", context_source="y", product_understanding_source_hash="z")
    with pytest.raises(ValidationError): UserStory(id="", title="", persona="", acceptance_criteria=[])
    assert ProductRequirementsDocument().risks is not ProductRequirementsDocument().risks

def test_parsers_handle_fences_and_failures():
    assert parse_prd_response("```json\n" + json.dumps(prd().model_dump()) + "\n```").title == "PRD"
    for bad in ("", "not json", "{}"):
        with pytest.raises(ArtifactParseError): parse_prd_response(bad)

def test_prompts_include_only_structured_understanding_and_metadata():
    for builder in (build_prd_prompt, build_user_stories_prompt, build_technical_specification_prompt):
        prompt = builder(understanding(), "Product Owner", "Validated Product Context")
        assert '"executive_summary": "Summary"' in prompt
        assert "Product Owner" in prompt and "Validated Product Context" in prompt
        assert "valid JSON only" in prompt and "raw uploaded documents" in prompt

def test_docx_all_artifacts_hide_hash_and_include_metadata():
    for content, kind in ((prd(), ArtifactType.PRD), (stories(), ArtifactType.USER_STORIES), (tech(), ArtifactType.TECHNICAL_SPECIFICATION)):
        data = build_artifact_docx(GeneratedArtifact(metadata=metadata(kind), content=content))
        text = "\n".join(paragraph.text for paragraph in Document(BytesIO(data)).paragraphs)
        assert data and content.title in text and kind.value in text and "secret-hash" not in text

def test_hash_staleness_detects_upstream_changes():
    current = build_product_understanding_source_hash("context", "Product Owner", "Original Product Context")
    assert not is_product_understanding_stale(current, "context", "Product Owner", "Original Product Context")
    assert is_product_understanding_stale(current, "context", "Technical Lead", "Original Product Context")
    assert is_product_understanding_stale(current, "changed", "Product Owner", "Original Product Context")
